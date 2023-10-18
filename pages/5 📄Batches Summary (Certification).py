import base64
from zipfile import ZipFile
import streamlit as st
import pandas as pd
import time
import json
from datetime import datetime
from docx import Document

st.set_page_config(
    page_title="Batches Summary (Certification)",
    page_icon="📄",
    layout="wide")

st.image("Images/Consolidation Tool.png")

if 'Xfile' not in st.session_state:
    st.session_state['Xfile']=None
    st.session_state['Xsheet']=None
    st.session_state['Dfile']=None

c1,c2=st.columns(2)
    

uploaded_files = c1.file_uploader('Upload Excel', type=['xlsx','xls','xlsb'])
st.session_state['Xfile']=uploaded_files

if st.session_state['Xfile']!=None:
    if uploaded_files.type!="text/csv":
        data=pd.ExcelFile(uploaded_files)
    else:
        data=pd.read_csv(uploaded_files)

    if uploaded_files.type!="text/csv":
        sheetmaster = c1.selectbox(
        'Select master sheet to proceed further !',
        (data.sheet_names))
        if c1.button("Load Master Sheet"):
            st.session_state['Xsheet']=sheetmaster
            
        if st.session_state['Xsheet']:
            data=data.parse(sheetmaster)
            c1.dataframe(data)

uploaded_doc = c2.file_uploader('Upload Document', type=['docx'])
st.session_state['Dfile']=uploaded_doc

if st.session_state['Xfile']!=None and st.session_state['Dfile']!=None:
    ZipfileDotZip = "Output files/SPHF -  Disbursement Certificate"+str(datetime.now().strftime("%Y-%m-%d %H-%M-%S "))+".zip"
    zipObj = ZipFile(ZipfileDotZip, "w")
    
    progress_text ="**** Kaam to hota rahy ga haal chaal kysy hai? ****"
    my_bar = st.progress(0, text=progress_text)
 
    master=data.copy()
    time.sleep(2)
    document = Document(uploaded_doc)

    def std(x):
        if type(x)==str:
            return x.split("(")[0].upper().split("/")[0].strip()
        else:
            return x
        
    master["DA_TypeSTD"]=master["DA_Type"].apply(std)

    # with pd.ExcelWriter("Output Batches Summary.xlsx") as writer:

    for batch in master["Batch#"].unique():
        bth=master[master["Batch#"]==batch]
        print(batch)
        
        document.paragraphs[6].text='Subject: '+str(batch)+': Housing Assistance Disbursement Statement '
        #document.paragraphs[6].style.font.bold=True
        
        document.tables[0].cell(1, 1).text=str(bth.shape[0])
        document.tables[0].cell(2, 1).text=str(bth[(bth["UUID_DEVIATION"]=="Not Matched")|(bth["UUID_DEVIATION"]=="Deviation")].shape[0])
        document.tables[0].cell(3, 1).text=str(bth[(bth["CNIC_DEVIATION"]=="Not Matched")|(bth["CNIC_DEVIATION"]=="Deviation")].shape[0])
        document.tables[0].cell(4, 1).text=str(bth[(bth["is_hazardous_location"]=="yes")].shape[0])
        document.tables[0].cell(5, 1).text=str(bth[(bth["Has_Reconstruction_Started"]=="yes")].shape[0])
        document.tables[0].cell(6, 1).text=str(bth[bth["landownership_type"].isin(
                                                                                ['nonServeyKachaLan',
                                                                                'multiOwner',
                                                                                'tenant'])].shape[0])
        
        document.paragraphs[21].text='We have found '+document.tables[0].cell(3, 1).text+' cases with CNIC deviation in this batch. As per SPHF policy (attached as Annexure-C), the caseload will be accepted in the following cases of the CNIC deviation:'
        document.paragraphs[18].text='We have found '+document.tables[0].cell(2, 1).text+' cases with JV&A ID deviation and all JV&A IDs in the damage assessment unique. '
        document.paragraphs[26].text='We have found '+document.tables[0].cell(4, 1).text+' cases with Hazardous location status in this batch. The caseload has been duly reviewed by SPHF team and based on the letter received from SPHF and Implementing Partners, the hazardous location risk has been revisited in light of guidance for resilient construction of houses. The supporting letters are annexed.   '
        document.paragraphs[28].text='We have found '+document.tables[0].cell(5, 1).text+' cases with reconstruction started, status in this batch. The caseload has been duly reviewed by SPHF technical team and technical guidance will be provided to ensure compliances against the minimum guidelines by SPHF. In case technical guidance is not followed, the cases shall not go forward for further disbursement. '
        document.paragraphs[30].text='We have found '+document.tables[0].cell(6, 1).text+' cases with land status of Non-surveyed Katcha, Tenant, and Multi-owner land from the list of beneficiaries provided.    '

        lst={"Total Cases in a Batch":bth.shape[0],
            "Cases with UU ID deviation":bth[(bth["UUID_DEVIATION"]=="Not Matched")|(bth["UUID_DEVIATION"]=="Deviation")].shape[0],
            "Cases CNIC Deviation":bth[(bth["CNIC_DEVIATION"]=="Not Matched")|(bth["CNIC_DEVIATION"]=="Deviation")].shape[0],
            "Cases in Hazardous Location":bth[bth["is_hazardous_location"]=="yes"].shape[0],
            "Cases with Reconstruction Started":bth[bth["Has_Reconstruction_Started"]=="yes"].shape[0],
            "Non-Surveyed Katcha / Tenant and Multi Owner Land status":bth[bth["landownership_type"].isin(
                                                                                ['nonServeyKachaLan',
                                                                                'multiOwner',
                                                                                'tenant'])].shape[0],

            }
    #     pd.DataFrame(lst.items(), columns=['Details', 'Number of cases']).to_excel(writer, sheet_name=batch,index=False)
        outputFileName='Output files/SPHF -  Disbursement Certificate - '+str(batch)+' - Update.docx'
        document.save(outputFileName)
        zipObj.write(outputFileName)
    zipObj.close()

    my_bar.progress(100, text="!!! chal bhai kaam hogaya aagy dekh ab  !!!")
    time.sleep(2)

    st.success("Processing Completed Click below to download output")

    with open(ZipfileDotZip, "rb") as f:
                bytes = f.read()
                b64 = base64.b64encode(bytes).decode()
                href = f"<a href=\"data:file/zip;base64,{b64}\" download='{ZipfileDotZip}.zip'>\
                    Click here to download the output\
                </a>"

    st.markdown(href, unsafe_allow_html=True)